"""Интеграционные тесты FastAPI: OCR и договор."""

import io
import pathlib
import sys
import types

import pytest
from fastapi.testclient import TestClient
from PIL import Image

_ROOT = pathlib.Path(__file__).resolve().parents[1]
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

if "pytesseract.pytesseract" not in sys.modules:
    _sub = types.ModuleType("pytesseract.pytesseract")
    _sub.TesseractError = type("TesseractError", (Exception,), {})
    sys.modules["pytesseract.pytesseract"] = _sub

if "pytesseract" not in sys.modules:
    _pkg = types.ModuleType("pytesseract")

    def _image_to_string(*_args, **_kwargs):
        return "stub ocr text"

    _pkg.image_to_string = _image_to_string
    _pkg.pytesseract = sys.modules["pytesseract.pytesseract"]
    sys.modules["pytesseract"] = _pkg

import dogovor
import main
import ocr
from pytesseract.pytesseract import TesseractError


@pytest.fixture()
def client():
    with TestClient(main.app) as c:
        yield c


@pytest.fixture()
def png_bytes() -> bytes:
    image = Image.new("RGB", (10, 10), color="white")
    buf = io.BytesIO()
    image.save(buf, format="PNG")
    return buf.getvalue()


@pytest.fixture()
def docx_template(tmp_path: pathlib.Path) -> pathlib.Path:
    """Минимальный шаблон договора для тестов без реального .docx."""
    from docx import Document

    path = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("Укажите ФИО")
    doc.add_paragraph("Укажите номер телефона")
    doc.save(path)
    return path


def test_root_returns_message(client):
    r = client.get("/")
    assert r.status_code == 200
    assert r.json() == {"message": "OCR API работает"}


def test_ocr_success(client, png_bytes):
    r = client.post(
        "/ocr",
        files={"file": ("ok.png", png_bytes, "image/png")},
    )
    assert r.status_code == 200
    body = r.json()
    assert body["filename"] == "ok.png"
    assert body["text"] == "stub ocr text"


def test_ocr_empty_file_returns_400(client):
    r = client.post(
        "/ocr",
        files={"file": ("empty.png", b"", "image/png")},
    )
    assert r.status_code == 400
    assert r.json()["detail"] == "Файл пустой"


def test_ocr_mp3_returns_415(client):
    r = client.post(
        "/ocr",
        files={"file": ("voice.mp3", b"fake-mp3", "audio/mpeg")},
    )
    assert r.status_code == 415
    assert "mp3" in r.json()["detail"].lower()


def test_ocr_too_large_returns_413(client, monkeypatch):
    monkeypatch.setattr(ocr, "MAX_FILE_SIZE_BYTES", 10)
    r = client.post(
        "/ocr",
        files={"file": ("big.png", b"0123456789", "image/png")},
    )
    assert r.status_code == 413
    assert "1 гб" in r.json()["detail"].lower()


def test_ocr_invalid_image_returns_422(client):
    r = client.post(
        "/ocr",
        files={"file": ("bad.bin", b"not-an-image-at-all", "application/octet-stream")},
    )
    assert r.status_code == 422
    assert "изображение" in r.json()["detail"].lower()


def test_ocr_tesseract_error_returns_500(client, png_bytes, monkeypatch):
    def _boom(_contents: bytes):
        raise TesseractError(1, "tesseract failed")

    monkeypatch.setattr(ocr, "img_ocr", _boom)
    r = client.post(
        "/ocr",
        files={"file": ("ok.png", png_bytes, "image/png")},
    )
    assert r.status_code == 500
    assert "Tesseract" in r.json()["detail"]


def test_extract_fields_finds_fio():
    text = "ФИО: Иванов Иван Иванович\nтелефон: +7 900 123-45-67"
    fields = dogovor.extract_fields(text)
    assert "Иванов" in fields["Укажите ФИО"]
    assert "+7" in fields["Укажите номер телефона"] or "900" in fields["Укажите номер телефона"]


def test_ocr_to_contract_success(client, png_bytes, docx_template, monkeypatch, tmp_path):
    out_dir = tmp_path / "generated"
    out_dir.mkdir()
    monkeypatch.setattr(dogovor, "TEMPLATE_PATH", docx_template)
    monkeypatch.setattr(dogovor, "OUTPUT_DIR", out_dir)

    r = client.post(
        "/ocr-to-contract",
        files={"file": ("scan.png", png_bytes, "image/png")},
    )
    assert r.status_code == 200, r.text
    data = r.json()
    assert data["message"] == "Договор успешно заполнен"
    assert data["filename"] == "scan.png"
    assert "generated_filename" in data
    assert "download_url" in data
    assert data["download_url"].startswith("/download/")

    name = data["generated_filename"]
    assert (out_dir / name).is_file()

    dl = client.get(f"/download/{name}")
    assert dl.status_code == 200
    assert dl.headers.get("content-type", "").startswith(
        "application/vnd.openxmlformats"
    ) or "wordprocessingml" in dl.headers.get("content-type", "")


def test_ocr_to_contract_empty_file_400(client):
    r = client.post(
        "/ocr-to-contract",
        files={"file": ("empty.png", b"", "image/png")},
    )
    assert r.status_code == 400


def test_ocr_to_contract_mp3_returns_415(client):
    r = client.post(
        "/ocr-to-contract",
        files={"file": ("voice.mp3", b"fake-mp3", "audio/mpeg")},
    )
    assert r.status_code == 415
    assert "mp3" in r.json()["detail"].lower()


def test_ocr_to_contract_too_large_returns_413(client, monkeypatch):
    monkeypatch.setattr(ocr, "MAX_FILE_SIZE_BYTES", 10)
    r = client.post(
        "/ocr-to-contract",
        files={"file": ("big.png", b"0123456789", "image/png")},
    )
    assert r.status_code == 413
    assert "1 гб" in r.json()["detail"].lower()


def test_ocr_to_contract_template_missing_returns_500(
    client, png_bytes, monkeypatch, tmp_path
):
    missing = tmp_path / "no_such_template.docx"
    monkeypatch.setattr(dogovor, "TEMPLATE_PATH", missing)
    monkeypatch.setattr(dogovor, "OUTPUT_DIR", tmp_path / "out")

    r = client.post(
        "/ocr-to-contract",
        files={"file": ("x.png", png_bytes, "image/png")},
    )
    assert r.status_code == 500
    assert "шаблон" in r.json()["detail"].lower()


def test_download_unknown_file_returns_404(client):
    # Имя в допустимом формате UUID.hex, файла нет в каталоге
    r = client.get("/download/filled_contract_00000000000000000000000000000000.docx")
    assert r.status_code == 404


def test_download_invalid_filename_returns_400(client):
    r = client.get("/download/not-a-valid-name.docx")
    assert r.status_code == 400


def test_fill_contract_replaces_placeholder(tmp_path, docx_template):
    out = tmp_path / "filled.docx"
    dogovor.fill_contract(
        str(docx_template),
        str(out),
        {"Укажите ФИО": "Петров Пётр"},
    )
    assert out.is_file()
    from docx import Document

    doc = Document(out)
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "Петров" in full
    assert "Укажите ФИО" not in full or "Петров" in full
