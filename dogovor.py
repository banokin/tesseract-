import re
import uuid
from pathlib import Path

from docx import Document
from fastapi import APIRouter, FastAPI, File, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from PIL import UnidentifiedImageError
from pytesseract.pytesseract import TesseractError

from concurrency import lifespan, run_blocking
from ocr import img_ocr, router as ocr_router, validate_upload

router = APIRouter(tags=["dogovor"])

_DOWNLOAD_NAME = re.compile(r"^filled_contract_[0-9a-f]{32}\.docx$")

BASE_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = BASE_DIR / "Новый договор.docx"
if not TEMPLATE_PATH.exists():
    _alt = BASE_DIR / "Новый договор.docx"
    if _alt.exists():
        TEMPLATE_PATH = _alt
OUTPUT_DIR = BASE_DIR / "generated_docs"
OUTPUT_DIR.mkdir(exist_ok=True)


def _resolve_download_path(filename: str) -> Path:
    """Безопасный путь к сгенерированному договору (без path traversal)."""
    safe_name = Path(filename).name
    if not _DOWNLOAD_NAME.fullmatch(safe_name):
        raise HTTPException(
            status_code=400,
            detail="Некорректное имя файла.",
        )
    out_root = OUTPUT_DIR.resolve()
    path = (out_root / safe_name).resolve()
    try:
        path.relative_to(out_root)
    except ValueError:
        raise HTTPException(
            status_code=400,
            detail="Некорректный путь к файлу.",
        )
    return path


def extract_fields(text: str) -> dict:
    """
    Пример извлечения полей из OCR текста.
    Регулярки потом подстроишь под реальный формат сканов.
    """

    def find(pattern: str, default: str = "") -> str:
        match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
        return match.group(1).strip() if match else default

    fields = {
        "Укажите номер договора": find(r"договор[а-я\s№]*[:\-]?\s*([A-Za-zА-Яа-я0-9\-\/]+)"),
        "Укажите дату договора": find(r"дата[а-я\s]*[:\-]?\s*(\d{2}\.\d{2}\.\d{4})"),
        "Укажите ФИО": find(r"фио[:\-]?\s*([А-ЯЁA-Z][А-ЯЁA-Zа-яёa-z\s\-]+)"),
        "Укажите адрес регистрации": find(r"адрес регистрации[:\-]?\s*(.+)"),
        "Укажите номер телефона": find(r"(?:телефон|тел\.?)[:\-]?\s*([\+\d\-\(\)\s]{6,})"),
        "Укажите адрес почты": find(r"(?:e-?mail|почта)[:\-]?\s*([^\s]+@[^\s]+)"),
        "укажите серию паспорта": find(r"серия[:\-]?\s*(\d{4})"),
        "укажите номер паспорта": find(r"номер[:\-]?\s*(\d{6})"),
        "Укажите орган, выдавший паспорт": find(r"выдан[:\-]?\s*(.+)"),
        "укажите дату выдачи паспорта": find(r"дата выдачи[:\-]?\s*(\d{2}\.\d{2}\.\d{4})"),
        "укажите код подразделения": find(r"код подразделения[:\-]?\s*([\d\-]+)"),
        "укажите место рождения": find(r"место рождения[:\-]?\s*(.+)"),
        "укажите дату рождения": find(r"дата рождения[:\-]?\s*(\d{2}\.\d{2}\.\d{4})"),
        "Выберите название объекта": find(r"объект[:\-]?\s*(.+)"),
        "Укажите назначение объекта": find(r"назначение[:\-]?\s*(.+)"),
        "Укажите площадь объекта м2": find(r"площадь[:\-]?\s*([\d\.,]+)"),
        "Укажите адрес объекта": find(r"адрес объекта[:\-]?\s*(.+)"),
        "Укажите кадастровый номер объекта": find(r"кадастровый номер[:\-]?\s*([0-9:\-]+)"),
        "Укажите документ-основание права собственности": find(
            r"основание права собственности[:\-]?\s*(.+)"
        ),
    }

    return fields


def replace_text_in_paragraph(paragraph, replacements: dict):
    for old_text, new_text in replacements.items():
        if old_text in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(old_text, new_text)


def replace_text_in_table(table, replacements: dict):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph, replacements)
            for nested_table in cell.tables:
                replace_text_in_table(nested_table, replacements)


def fill_contract(template_path: str, output_path: str, fields: dict):
    doc = Document(template_path)
    replacements = {k: v for k, v in fields.items() if v}

    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    for table in doc.tables:
        replace_text_in_table(table, replacements)

    doc.save(output_path)


@router.post("/ocr-to-contract")
async def ocr_to_contract(request: Request, file: UploadFile = File(...)):
    try:
        contents = await file.read()
        validate_upload(file, contents)

        text = await run_blocking(request, img_ocr, contents)
        fields = await run_blocking(request, extract_fields, text)

        template_file = TEMPLATE_PATH
        if not template_file.exists():
            raise HTTPException(status_code=500, detail="Файл шаблона договора не найден")

        output_name = f"filled_contract_{uuid.uuid4().hex}.docx"
        output_path = OUTPUT_DIR / output_name

        await run_blocking(
            request,
            fill_contract,
            str(template_file),
            str(output_path),
            fields,
        )

        return {
            "message": "Договор успешно заполнен",
            "filename": file.filename,
            "generated_filename": output_name,
            "ocr_text": text,
            "fields": fields,
            "download_url": f"/download/{output_name}",
        }

    except UnidentifiedImageError:
        raise HTTPException(
            status_code=422,
            detail="Не удалось распознать изображение. Проверьте формат файла.",
        )
    except TesseractError as e:
        raise HTTPException(status_code=500, detail=f"Ошибка Tesseract: {str(e)}")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Ошибка обработки документа: {str(e)}")


@router.get("/download/{filename}")
async def download_file(filename: str):
    file_path = _resolve_download_path(filename)

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Файл не найден")

    return FileResponse(
        path=file_path,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _create_standalone_app() -> FastAPI:
    """Для запуска: uvicorn dogovor:app (OCR + договор)."""
    application = FastAPI(lifespan=lifespan)
    application.add_middleware(
        CORSMiddleware,
        allow_origins=["*"],
        allow_credentials=True,
        allow_methods=["*"],
        allow_headers=["*"],
    )

    @application.get("/")
    def root():
        return {"message": "OCR API работает"}

    application.include_router(ocr_router)
    application.include_router(router)
    return application


app = _create_standalone_app()
